"""
Service for synchronizing database with offers folder
"""
import os
import sqlite3
import json
from docx import Document
from typing import List, Dict, Tuple, Optional
import sys

# Add project root to Python path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(__file__))))

from src.utils.config import DATABASE_PATH, get_offers_folder
from src.data.database_service import get_client_by_nip


class SyncResult:
    """Class to hold synchronization results"""
    
    def __init__(self):
        self.success = True
        self.errors = []
        self.warnings = []
        self.total_offers_db = 0
        self.total_offers_folder = 0
        self.missing_files = []
        self.orphaned_files = []
        self.invalid_offer_numbers = []
        self.invalid_client_aliases = []


class OfferSyncService:
    """Service for synchronizing offers database with folder"""
    
    def __init__(self):
        self.offers_folder = get_offers_folder()
        self.db_path = DATABASE_PATH
    
    def synchronize(self) -> SyncResult:
        """Perform full synchronization between database and folder"""
        result = SyncResult()
        
        try:
            # First check if required database tables exist
            self._check_database_tables()
            
            # Step 1: Get all offers from database
            db_offers = self._get_database_offers()
            result.total_offers_db = len(db_offers)
            
            # Step 2: Get all offer files from folder
            folder_files = self._get_folder_offer_files()
            result.total_offers_folder = len(folder_files)
            
            # Step 3: Check each database offer
            for offer in db_offers:
                self._check_database_offer(offer, result)
            
            # Step 4: Find orphaned files (files without database entries)
            self._find_orphaned_files(db_offers, folder_files, result)
            
            # Step 5: Determine overall result
            if result.errors or result.missing_files or result.orphaned_files or result.invalid_offer_numbers or result.invalid_client_aliases:
                result.success = False
            
        except Exception as e:
            result.success = False
            result.errors.append(f"Błąd podczas synchronizacji: {str(e)}")
        
        return result
    
    def _check_database_tables(self):
        """Check if all required database tables exist"""
        required_tables = ['Offers', 'Clients', 'Suppliers']
        
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            for table in required_tables:
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name=?", (table,))
                if not cursor.fetchone():
                    conn.close()
                    raise Exception(f"Tabela '{table}' nie istnieje w bazie danych. Baza danych nie została zainicjalizowana lub jest uszkodzona.")
            
            conn.close()
            
        except sqlite3.Error as e:
            raise Exception(f"Błąd dostępu do bazy danych: {str(e)}")
    
    def _get_database_offers(self) -> List[Dict]:
        """Get all offers from database"""
        offers = []
        
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute("SELECT OfferFilePath, OfferContext FROM Offers")
            rows = cursor.fetchall()
            
            for file_path, context_json in rows:
                try:
                    context = json.loads(context_json) if context_json else {}
                    offers.append({
                        'file_path': file_path,
                        'context': context
                    })
                except json.JSONDecodeError:
                    offers.append({
                        'file_path': file_path,
                        'context': {}
                    })
            
            conn.close()
            
        except Exception as e:
            print(f"Error getting database offers: {e}")
            # Re-raise the exception so it can be caught by synchronize()
            raise
        
        return offers
    
    def _get_folder_offer_files(self) -> List[str]:
        """Get all .docx files from offers folder"""
        files = []
        
        try:
            if os.path.exists(self.offers_folder):
                # Include year subfolders
                files = []
                for entry in os.listdir(self.offers_folder):
                    year_dir = os.path.join(self.offers_folder, entry)
                    if os.path.isdir(year_dir) and entry.isdigit() and len(entry) == 4:
                        for f in os.listdir(year_dir):
                            if f.endswith('.docx'):
                                files.append(f"{entry}/{f}")
                    else:
                        if entry.endswith('.docx'):
                            files.append(entry)
        except Exception as e:
            print(f"Error reading offers folder: {e}")
        
        return files
    
    def _check_database_offer(self, offer: Dict, result: SyncResult):
        """Check single database offer against folder"""
        file_path = offer['file_path']
        context = offer['context']
        
        # Extract filename from path
        # file_path from DB may be relative 'YYYY/filename' or filename only
        parts = (file_path or '').replace('\\', '/').split('/')
        if len(parts) == 2:
            full_path = os.path.join(self.offers_folder, parts[0], parts[1])
            filename = parts[1]
        else:
            filename = parts[-1]
            full_path = os.path.join(self.offers_folder, filename)
        
        # Check if file exists
        if not os.path.exists(full_path):
            result.missing_files.append({
                'filename': filename,
                'db_path': file_path,
                'reason': 'Plik nie istnieje w folderze'
            })
            return
        
        # Check offer number consistency
        if not self._validate_offer_number(full_path, filename, result):
            return
        
        # Check client alias consistency
        if not self._validate_client_alias(full_path, filename, context, result):
            return
    
    def _validate_offer_number(self, file_path: str, filename: str, result: SyncResult) -> bool:
        """Validate offer number consistency between Word doc and filename"""
        try:
            # Extract offer number from Word document
            doc_offer_number = self._extract_offer_number_from_word(file_path)
            if not doc_offer_number:
                result.invalid_offer_numbers.append({
                    'filename': filename,
                    'reason': 'Nie można odczytać numeru oferty z dokumentu Word'
                })
                return False
            
            # Extract expected filename from offer number (replace / with _)
            expected_filename_base = doc_offer_number.replace('/', '_')
            
            # Get actual filename without extension
            actual_filename_base = os.path.splitext(filename)[0]
            
            # Compare
            if expected_filename_base != actual_filename_base:
                result.invalid_offer_numbers.append({
                    'filename': filename,
                    'doc_offer_number': doc_offer_number,
                    'expected_filename': f"{expected_filename_base}.docx",
                    'reason': 'Numer oferty w dokumencie nie odpowiada nazwie pliku'
                })
                return False
                
            return True
            
        except Exception as e:
            result.invalid_offer_numbers.append({
                'filename': filename,
                'reason': f'Błąd podczas walidacji numeru oferty: {str(e)}'
            })
            return False
    
    def _validate_client_alias(self, file_path: str, filename: str, context: Dict, result: SyncResult) -> bool:
        """Validate client alias consistency"""
        try:
            # Extract client NIP from Word document
            client_nip = self._extract_client_nip_from_word(file_path)
            if not client_nip:
                result.invalid_client_aliases.append({
                    'filename': filename,
                    'reason': 'Nie można odczytać NIP klienta z dokumentu Word'
                })
                return False
            
            # Get client alias from filename
            filename_parts = os.path.splitext(filename)[0].split('_')
            if len(filename_parts) < 4:
                result.invalid_client_aliases.append({
                    'filename': filename,
                    'reason': 'Nieprawidłowy format nazwy pliku - brak aliasu klienta'
                })
                return False
            
            file_client_alias = filename_parts[-1]  # Last part should be client alias
            
            # Find client in database by NIP
            client_data = get_client_by_nip(client_nip)
            if not client_data:
                result.invalid_client_aliases.append({
                    'filename': filename,
                    'client_nip': client_nip,
                    'reason': 'Nie znaleziono klienta w bazie danych po NIP'
                })
                return False
            
            # client_data format: (nip, company_name, address1, address2, alias)
            db_client_alias = client_data[4]
            
            if file_client_alias != db_client_alias:
                result.invalid_client_aliases.append({
                    'filename': filename,
                    'file_alias': file_client_alias,
                    'db_alias': db_client_alias,
                    'client_nip': client_nip,
                    'reason': 'Alias klienta w nazwie pliku nie odpowiada aliasowi w bazie danych'
                })
                return False
            
            return True
            
        except Exception as e:
            result.invalid_client_aliases.append({
                'filename': filename,
                'reason': f'Błąd podczas walidacji aliasu klienta: {str(e)}'
            })
            return False
    
    def _extract_offer_number_from_word(self, file_path: str) -> Optional[str]:
        """Extract offer number from Word document"""
        try:
            doc = Document(file_path)
            
            # Look for offer number in document text
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                
                # Look for various patterns
                patterns = [
                    'OFERTA NR:',
                    'Nr oferty:',
                    'Numer oferty:',
                    'OFFER NUMBER:',
                    'OFERTA:'
                ]
                
                for pattern in patterns:
                    if pattern in text:
                        # Extract number after the pattern
                        parts = text.split(pattern)
                        if len(parts) > 1:
                            offer_number = parts[1].strip()
                            return offer_number
            
            # Alternative: look in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        patterns = [
                            'OFERTA NR:',
                            'Nr oferty:',
                            'Numer oferty:',
                            'OFFER NUMBER:',
                            'OFERTA:'
                        ]
                        
                        for pattern in patterns:
                            if pattern in text:
                                parts = text.split(pattern)
                                if len(parts) > 1:
                                    offer_number = parts[1].strip()
                                    return offer_number
            
            return None
            
        except Exception as e:
            print(f"Error extracting offer number from {file_path}: {e}")
            return None
    
    def _extract_client_nip_from_word(self, file_path: str) -> Optional[str]:
        """Extract client NIP from Word document"""
        try:
            doc = Document(file_path)
            
            # Strategy 1: Look for NIP after "KLIENT:" section
            klient_found = False
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                
                if 'KLIENT:' in text.upper():
                    klient_found = True
                    continue
                
                # If we found KLIENT section, look for next NIP
                if klient_found and 'NIP:' in text.upper():
                    parts = text.split(':')
                    if len(parts) > 1:
                        nip = parts[1].strip()
                        if nip:
                            return self._clean_nip(nip)
            
            # Strategy 2: Look for multiple NIPs and take the last one (likely client)
            nips_found = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if 'NIP:' in text.upper():
                    parts = text.split(':')
                    if len(parts) > 1:
                        nip = parts[1].strip()
                        if nip:
                            nips_found.append(self._clean_nip(nip))
            
            # Return the last NIP found (likely the client's)
            if len(nips_found) > 1:
                return nips_found[-1]
            elif len(nips_found) == 1:
                return nips_found[0]
            
            # Strategy 3: Look in tables
            for table in doc.tables:
                for row in table.rows:
                    for i, cell in enumerate(row.cells):
                        text = cell.text.strip()
                        if 'NIP:' in text.upper():
                            # Try to get NIP from same cell or next cell
                            if ':' in text:
                                parts = text.split(':')
                                if len(parts) > 1:
                                    nip = parts[1].strip()
                                    if nip:
                                        return self._clean_nip(nip)
                            
                            # Try next cell
                            if i + 1 < len(row.cells):
                                nip = row.cells[i + 1].text.strip()
                                if nip:
                                    return self._clean_nip(nip)
            
            return None
            
        except Exception as e:
            print(f"Error extracting client NIP from {file_path}: {e}")
            return None
    
    def _clean_nip(self, nip: str) -> str:
        """Clean NIP string to contain only digits"""
        return ''.join(c for c in nip if c.isdigit())
    
    def _find_orphaned_files(self, db_offers: List[Dict], folder_files: List[str], result: SyncResult):
        """Find files in folder that don't have database entries"""
        db_filenames = set()
        
        for offer in db_offers:
            filename = os.path.basename(offer['file_path'])
            db_filenames.add(filename)
        
        for file in folder_files:
            if file not in db_filenames:
                result.orphaned_files.append({
                    'filename': file,
                    'reason': 'Plik nie ma wpisu w bazie danych'
                })
    
    def generate_report(self, result: SyncResult) -> str:
        """Generate detailed synchronization report"""
        report = []
        report.append("=== RAPORT SYNCHRONIZACJI BAZY DANYCH Z FOLDEREM OFERT ===\n")
        
        if result.success:
            report.append("✅ SYNCHRONIZACJA ZAKOŃCZONA POMYŚLNIE")
            report.append(f"Sprawdzono {result.total_offers_db} ofert z bazy danych")
            report.append(f"Znaleziono {result.total_offers_folder} plików w folderze ofert")
            report.append("Wszystkie dane są spójne.\n")
        else:
            report.append("❌ WYKRYTO PROBLEMY PODCZAS SYNCHRONIZACJI")
            report.append(f"Sprawdzono {result.total_offers_db} ofert z bazy danych")
            report.append(f"Znaleziono {result.total_offers_folder} plików w folderze ofert\n")
            
            if result.errors:
                report.append("🔴 BŁĘDY OGÓLNE:")
                for error in result.errors:
                    report.append(f"  - {error}")
                report.append("")
            
            if result.missing_files:
                report.append("📁 BRAKUJĄCE PLIKI (wpisy w bazie, ale brak plików):")
                for item in result.missing_files:
                    report.append(f"  - {item['filename']}: {item['reason']}")
                report.append("")
            
            if result.orphaned_files:
                report.append("🗂️ PLIKI SIEROTY (pliki bez wpisów w bazie):")
                for item in result.orphaned_files:
                    report.append(f"  - {item['filename']}: {item['reason']}")
                report.append("")
            
            if result.invalid_offer_numbers:
                report.append("🔢 NIEPRAWIDŁOWE NUMERY OFERT:")
                for item in result.invalid_offer_numbers:
                    if 'doc_offer_number' in item:
                        report.append(f"  - {item['filename']}")
                        report.append(f"    Nr z dokumentu: {item['doc_offer_number']}")
                        report.append(f"    Oczekiwana nazwa: {item['expected_filename']}")
                        report.append(f"    Problem: {item['reason']}")
                    else:
                        report.append(f"  - {item['filename']}: {item['reason']}")
                report.append("")
            
            if result.invalid_client_aliases:
                report.append("👤 NIEPRAWIDŁOWE ALIASY KLIENTÓW:")
                for item in result.invalid_client_aliases:
                    if 'file_alias' in item:
                        report.append(f"  - {item['filename']}")
                        report.append(f"    Alias w pliku: {item['file_alias']}")
                        report.append(f"    Alias w bazie: {item['db_alias']}")
                        report.append(f"    NIP klienta: {item['client_nip']}")
                        report.append(f"    Problem: {item['reason']}")
                    else:
                        report.append(f"  - {item['filename']}: {item['reason']}")
                report.append("")
        
        return "\n".join(report)
